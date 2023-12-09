import openai
import json
from docx import Document
from ..models import ProjectRequirementDocument
def generate_prd( project_description, project_timeline, project_techstacks):

    # Define the prompt based on parameters
    prompt = f"""
        You Are a Project Manager. Generate a comprehensive Product Requirements Document (PRD) for a new project. Fill In Every Minute Detail Present.
        Provide the project details in a JSON data structure, ensuring that each key corresponds to the following elements:
        The project involves {project_description}. Below are the key details:


        ### Project Description:
        {project_description}

        ### Project Timeline:
        {project_timeline}

        ### Project Techstacks
        {project_techstacks}

        Considering this information, generate a PRD that adheres to the following structure:

        1. **Project Overview:**
            [Provide a concise overview of the project, including its purpose, scope, and key features.]

        2. **Original Requirements:**
            [Specify the original requirements based on the project's needs. Outline both functional and non-functional requirements.]

        3. **Project Goals:**
            [Define up to 3 clear and orthogonal project goals. Align them with the overall vision and success criteria of the project.]

        4. **User Stories:**
            [Present up to 5 scenario-based user stories. Capture diverse use cases, user interactions, and personas.]

        5. **System Architecture:**
            [Outline the high-level system architecture, covering both hardware and software components. Describe how they interact to meet project goals.]

        6. **Tech Stacks:**
            [Taking the project description into consideration, list all the tech stacks that will be used in the project.]

        7. **Requirement Pool:**
            [List up to 5 key requirements with priority (P0/P1/P2) and brief descriptions. Align each requirement with project goals.]

        8. **UI/UX Design:**
            [Provide a detailed plain-text description of the UI/UX design. Include elements, functions, style, and layout details.]

        9. **Development Methodology:**
            [Specify the development methodology (e.g., Agile, Waterfall) and explain how development phases, testing, and deployment will be managed.]

        10. **Security Measures:**
            [Detail security measures for both hardware and software components. Discuss encryption, access controls, and measures to protect user data.]

        11. **Testing Strategy:**
            [Describe the testing strategy, including types of testing (e.g., unit, integration) for both hardware and software components.]

        12. **Scalability and Performance:**
            [Address scalability and performance considerations for both hardware and software. Discuss how the system will handle increased load.]

        13. **Deployment Plan:**
            [Outline the deployment plan, specifying steps for deploying software updates and managing hardware deployment.]

        14. **Maintenance and Support:**
            [Define the plan for ongoing maintenance and support, including issue resolution and updates for both hardware and software.]

        15. **Risks and Mitigations:**
            [Identify potential risks associated with the project, proposing mitigation strategies for both hardware and software aspects.]

        16. **Compliance and Regulations:**
            [Ensure the project complies with relevant regulations and standards. Outline any certifications or compliance measures required.]

        17. **Budget and Resources:**
            [Provide an overview of the budget and resources allocated, covering both hardware and software development.]

        18. **Timeline and Milestones:**
            [Outline the project timeline and key milestones, considering both hardware and software development phases.]

        19. **Communication Plan:**
            [Define a communication plan for stakeholders, ensuring clear and effective communication throughout the project.]

        20. **Anything UNCLEAR:**
            [Address any uncertainties or unclear points in the project. Provide clarifications or assumptions. Encourage further questions or discussions.]

       Ensure that the generated PRD is clear, concise, and provides all necessary information for stakeholders to understand and proceed with the project.
        Please provide the information in only in JSON format for easy storage in a database.
        JSON format is highly important for the response to be successful.
        JSON format example:
        """
    prompt+="""{
            "name": "sonoo",
            "salary": 56000,
            "married": true,
        }"""

    openai.api_key = "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"
    response = openai.Completion.create(
        engine="gpt-3.5-turbo-instruct",
        prompt=prompt,
        max_tokens=3000,  # Adjust as needed
        temperature=0.7,  # Adjust as needed
    )

    return response.choices[0].text.strip()


def create_word_document(content, filename="output.docx"):
    # Create a Word document
    doc = Document()
    doc.add_heading("Generated PRD", level=1)
    doc.add_paragraph(content)

    # Save the Word document
    doc.save(filename)
    print(f"Document saved as '{filename}'")


# When generate PRD button is clicked
def generate_prd_button_clicked(project):
    # Get all the values from the database
    project_techstacks = project.related_techstacks
    project_description = project.description
    project_timeline = f"{project.end_date} to {project.start_date}"

    # Generate PRD
    prd = generate_prd(project_description, project_timeline, project_techstacks)

    # Create a Word document
    create_word_document(prd)
    print(prd)
    json_response = json.loads(prd, strict=False)
    try:
        project_details = ProjectRequirementDocument(
            project_overview=json_response["Project Overview"],
            original_requirements=json_response["Original Requirements"],
            project_goals=json_response["Project Goals"],
            user_stories=json_response["User Stories"],
            system_architecture=json_response["System Architecture"],
            tech_stacks=json_response["Tech Stacks"],
            requirement_pool=json_response["Requirement Pool"],
            ui_ux_design=json_response["UI/UX Design"],
            development_methodology=json_response["Development Methodology"],
            security_measures=json_response["Security Measures"],
            testing_strategy=json_response["Testing Strategy"],
            scalability_and_performance=json_response["Scalability and Performance"],
            deployment_plan=json_response["Deployment Plan"],
            maintenance_and_support=json_response["Maintenance and Support"],
            risks_and_mitigations=json_response["Risks and Mitigations"],
            compliance_and_regulations=json_response["Compliance and Regulations"],
            budget_and_resources=json_response["Budget and Resources"],
            timeline_and_milestones=json_response["Timeline and Milestones"],
            communication_plan=json_response["Communication Plan"],
            anything_unclear=json_response["Anything UNCLEAR"],
        )
    except KeyError:
        project_details = ProjectRequirementDocument(
            project_overview=json_response["project_overview"],
            original_requirements=json_response["original_requirements"],
            project_goals=json_response["project_goals"],
            user_stories=json_response["user_stories"],
            system_architecture=json_response["system_architecture"],
            tech_stacks=json_response["tech_stacks"],
            requirement_pool=json_response["requirement_pool"],
            ui_ux_design=json_response["ui_ux_design"],
            development_methodology=json_response["development_methodology"],
            security_measures=json_response["security_measures"],
            testing_strategy=json_response["testing_strategy"],
            scalability_and_performance=json_response["scalability_and_performance"],
            deployment_plan=json_response["deployment_plan"],
            maintenance_and_support=json_response["maintenance_and_support"],
            risks_and_mitigations=json_response["risks_and_mitigations"],
            compliance_and_regulations=json_response["compliance_and_regulations"],
            budget_and_resources=json_response["budget_and_resources"],
            timeline_and_milestones=json_response["timeline_and_milestones"],
            communication_plan=json_response["communication_plan"],
            anything_unclear=json_response["anything_unclear"],
        )
    project_details.save()
    project.prd = project_details
    project.save()
    return project_details.id

