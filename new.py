import openai
from docx import Document

class MessageHandlerForClient:
    def __init__(self, api_key):
        self.roles = "Alumni: Project Owner, Project Owner: Project Manager, Project Manager: Technical Lead, Technical Lead: Developer, Developer: QA Engineer, QA Engineer: Project Manager"
        openai.api_key = api_key
        self.role = "Alumni"

    

    def __is__question__answerable__(self, question):
        response = openai.Completion.create(
            engine="davinci",
            prompt=f"Is this question answerable by any of the above roles {self.roles}? The question is as follows: {question} .Just Return Yes Or No.",
            max_tokens=50,
        )

        answer = response.choices[0].text.strip()
        answer = answer.toLower()
        return answer == "yes"
    def decide_role(self, question):
        if (self.__is__question__answerable__(question)):
            #if the question is answerable by any of the above roles , then it goes into this for loop
            response = openai.Completion.create(
                engine="davinci",
                prompt=f"Which role among {self.roles} should handle this question? The question is as follows: {question}. Give your response in just the name of the role.",
                max_tokens=50,
            )

            role = response.choices[0].text.strip()
            return role
        else:
            #if the question is not answerable by any of the above roles , then it goes into this for loop , here AI answers it on its own
            response = openai.Completion.create(
                engine="davinci",
                prompt=f"Which role among {self.roles} should handle this question? The question is as follows: {question}. Give your response in just the name of the role.",
                max_tokens=50,
            )

            role = response.choices[0].text.strip()
            return role

    def handle_message(self, user_question, project_requirements, project_description, project_timeline, project_milestones):
        role = self.decide_role(user_question)
        
        response = openai.Completion.create(
            engine="davinci", prompt=f"As a {role}, {user_question}. Give your response as you are best in that field and you really want to help the user .", max_tokens=50
        )
        answer = response.choices[0].text.strip()
        return answer
    def handle_project_related_messages(self, user_question, project_requirements, project_description, project_timeline, project_milestones,students_skills):
        students_info = ", ".join([f'{student}: {skills}' for student, skills in students_skills.items()])
        project_details = f"""
            The project involves {project_description}. Below are the key details:

            ### Project Requirements:
            {project_requirements}


            ### Project Description:
            {project_description}

            ### Project Timeline:
            {project_timeline}

            ### Project Milestones:
            {project_milestones}

            ### Student Skills:
            {students_info}
        """
        response = openai.Completion.create(
            engine="davinci", prompt=f"As a HR manager or General Advisor,You have a Question That You Need To Answer A Client {user_question}.{project_details} , the above details may help you in answering the question Give your response as you are best in that field and you really want to help the user . just return the answer and nothing else . ", max_tokens=50
        )
        answer = response.choices[0].text.strip()
        return answer
    def generate_prd(self,project_requirements, project_description, project_timeline, project_milestones):
        # Define the prompt based on parameters
        
        prompt = f"""
        You Are a Project Manager. Generate a comprehensive Product Requirements Document (PRD) for a new project. Fill In Every Minute Detail Present , Dont Give Any Code and Make it Compatible For A Word File
        The project involves {project_description}. Below are the key details:

        ### Project Requirements:
        {project_requirements}


        ### Project Description:
        {project_description}

        ### Project Timeline:
        {project_timeline}

        ### Project Milestones:
        {project_milestones}

        Considering this information, generate a PRD that adheres to the following structure:

        1. **Project Overview:**
            [Provide a concise overview of the project, including its purpose, scope, and key features.]

        2. **Original Requirements:**
            [Specify the original requirements based on the project's needs. Outline both functional and non-functional requirements.]

        3. **Project Goals:**
            [Define up to 3 clear and orthogonal project goals. Align them with the overall vision and success criteria of the project.]

        4. **User Stories:**
            [Present up to 5 scenario-based user stories. Capture diverse use cases, user interactions, and personas.]

        5. **System Architecture:**
            [Outline the high-level system architecture, covering both hardware and software components. Describe how they interact to meet project goals.]

        6. **Requirements Analysis:**
            and their impact on project success.]

        7. **Requirement Pool:**
            [List up to 5 key requirements with priority (P0/P1/P2) and brief descriptions. Align each requirement with project goals.]

        8. **UI/UX Design:**
            [Provide a detailed plain-text description of the UI/UX design. Include elements, functions, style, and layout details.]

        9. **Development Methodology:**
            [Specify the development methodology (e.g., Agile, Waterfall) and explain how development phases, testing, and deployment will be managed.]

        10. **Security Measures:**
            [Detail security measures for both hardware and software components. Discuss encryption, access controls, and measures to protect user data.]

        11. **Testing Strategy:**
            [Describe the testing strategy, including types of testing (e.g., unit, integration) for both hardware and software components.]

        12. **Scalability and Performance:**
            [Address scalability and performance considerations for both hardware and software. Discuss how the system will handle increased load.]

        13. **Deployment Plan:**
            [Outline the deployment plan, specifying steps for deploying software updates and managing hardware deployment.]

        14. **Maintenance and Support:**
            [Define the plan for ongoing maintenance and support, including issue resolution and updates for both hardware and software.]

        15. **Risks and Mitigations:**
            [Identify potential risks associated with the project, proposing mitigation strategies for both hardware and software aspects.]

        16. **Compliance and Regulations:**
            [Ensure the project complies with relevant regulations and standards. Outline any certifications or compliance measures required.]

        17. **Budget and Resources:**
            [Provide an overview of the budget and resources allocated, covering both hardware and software development.]

        18. **Timeline and Milestones:**
            [Outline the project timeline and key milestones, considering both hardware and software development phases.]

        19. **Communication Plan:**
            [Define a communication plan for stakeholders, ensuring clear and effective communication throughout the project.]

        20. **Anything UNCLEAR:**
            [Address any uncertainties or unclear points in the project. Provide clarifications or assumptions. Encourage further questions or discussions.]

        Feel free to customize this template based on the specific needs and nature of your project.

        Ensure that the generated PRD is clear, concise, and provides all necessary information for stakeholders to understand and proceed with the project.
        """

        # Use OpenAI to generate PRD based on the prompt
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=1000,  # Adjust as needed
            temperature=0.7,  # Adjust as needed
        )

        return response.choices[0].text.strip()

    def create_word_document(self,content, filename="output.docx"):
        # Create a Word document
        doc = Document()
        doc.add_heading('Generated PRD', level=1)
        doc.add_paragraph(content)

        # Save the Word document
        doc.save(filename)
        print(f"Document saved as '{filename}'")
    def generate_project_workflow_prompt(self,project_description, project_requirements, project_timeline, project_milestones, students_skills):
        # Joining student information into a string
        students_info = ", ".join([f'{student}: {skills}' for student, skills in students_skills.items()])

        prompt = f"""
            As a Project Manager, your task is to generate a comprehensive workflow for a new project based on the skills of the assigned students. The project involves {project_description}. Below are the key details:

            ### Project Requirements:
            {project_requirements}

            ### Project Description:
            {project_description}

            ### Project Timeline:
            {project_timeline}

            ### Project Milestones:
            {project_milestones}

            ### Student Skills:
            The following students have been assigned to the project along with their respective skillsets:
            
            {students_info}

            ### Workflow Integration:
            Considering the skills of each student, outline a detailed workflow that leverages their expertise. Ensure efficient task allocation and collaboration among the team members.

            Note: Do not provide actual code; instead, create a narrative or bullet-point format suitable for a Word file.
        """
        return prompt
    def make_workflow(self,project_requirements, project_description, project_timeline, project_milestones,students_skills):
        prompt = self.generate_project_workflow_prompt(project_description, project_requirements, project_timeline, project_milestones, students_skills)
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=1000,  # Adjust as needed
            temperature=0.7,  # Adjust as needed
        )

        return response.choices[0].text.strip()
        
        
api_key = "YOUR_OPENAI_API_KEY"  # Replace with your actual OpenAI API key
handler = MessageHandlerForClient(api_key)

question = "Sample question template"
print(handler.handle_message(question))
